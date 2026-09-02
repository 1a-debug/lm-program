from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parents[1] / "lm-program五分钟完整演示演讲稿.docx"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "203040"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GREEN = "18794E"
GOLD = "8A6500"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_command(doc, text: str):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, "Consolas", 9.5, INK)
    return p


def add_cue(doc, label: str, text: str, color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=110, bottom=110, start=150, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(label + "  ")
    set_run_font(label_run, size=10, color=color, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10, color=INK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_speech(doc, text: str):
    p = doc.add_paragraph(style="Speech")
    p.paragraph_format.keep_together = True
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def add_step(doc, time: str, title: str):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.keep_with_next = True
    time_run = p.add_run(time + "  ")
    set_run_font(time_run, size=13, color=GOLD, bold=True)
    title_run = p.add_run(title)
    set_run_font(title_run, size=13, color=BLUE, bold=True)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    title = styles["Title"]
    title.font.name = "Microsoft YaHei"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(27)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    title.paragraph_format.space_after = Pt(8)

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK_BLUE),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(9.5)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.05
    p_pr = code.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_BLUE)
    p_pr.append(shd)

    speech = styles.add_style("Speech", 1)
    speech.font.name = "Microsoft YaHei"
    speech.font.size = Pt(10.5)
    speech.paragraph_format.left_indent = Inches(0.16)
    speech.paragraph_format.space_after = Pt(7)
    speech.paragraph_format.line_spacing = 1.25


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.text = "LM-PROGRAM  |  五分钟演示"
    set_run_font(header.runs[0], size=8.5, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("演示讲稿  ·  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(footer, "PAGE")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    kicker = p.add_run("CODING AGENT DEMO SCRIPT")
    set_run_font(kicker, size=10, color=GOLD, bold=True)
    title = doc.add_paragraph("lm-program 五分钟完整演示演讲稿", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = doc.add_paragraph("先展示基础编程闭环，再展示 Agent Guardian 与 Blocker Gate")
    subtitle.paragraph_format.space_after = Pt(22)
    set_run_font(subtitle.runs[0], size=13, color=MUTED)

    add_cue(doc, "演示目标", "在 5 分钟内证明它是真正执行本地工具的 Coding Agent，并突出测试完整性、提示注入防护和无效探索阻断。", GREEN)
    add_cue(doc, "必须保留的证据", "真实测试失败、edit 差异、修复后 OK、Guardian 3/3 PASS、SHA-256 MATCH、Blocker Gate ACTIVATED。", GOLD)

    doc.add_heading("录制前准备", level=1)
    add_command(doc, "cd D:\\coding-agent-main\nuv run lm-program --cwd demo\\calculator-project")
    add_speech(doc, "确认启动画面出现 features: Guardian + Blocker Gate，并确保终端中没有显示 API Key。看到“>”才是 Agent 输入框；看到“PS D:\\...”说明已经退出 Agent。")

    doc.add_page_break()
    doc.add_heading("第一部分  基础 Coding Agent 能力", level=1)
    add_step(doc, "0:00–0:30", "开场介绍")
    add_speech(doc, "大家好，我的项目是 lm-program，一个使用 Python 实现的终端 Coding Agent。用户输入编程任务后，大语言模型可以调用本地工具读取文件、修改代码和执行命令，程序再把执行结果返回给模型，让模型根据真实结果继续分析。下面我先演示基础编程能力，再展示项目的两个特色功能。")

    add_step(doc, "0:30–0:45", "恢复可重复的 Bug")
    add_command(doc, "/demo-reset")
    add_speech(doc, "我先通过 demo-reset 恢复一个预设 Bug。这个命令只恢复错误的实现，不修改测试文件，所以后面的失败和修复都是真实发生的，并且可以重复演示。")

    add_step(doc, "0:45–2:15", "让 Agent 自主定位并修复")
    add_command(doc, "运行全部测试，定位失败原因，只修改实现代码，不要修改测试，修复后重新运行测试验证。")
    add_cue(doc, "看到 read_file 时", "说明模型正在读取 calculator.py 和 test_calculator.py，而不是直接输出固定答案。")
    add_speech(doc, "Agent 首先读取实现和测试代码，然后选择测试命令。它需要根据项目中的真实文件决定下一步操作。")
    add_cue(doc, "看到 pytest 未安装时", "No module named pytest 不是项目 Bug；观察 Agent 是否主动改用 unittest。", GOLD)
    add_speech(doc, "模型最初尝试使用 pytest，但当前环境没有安装。Agent 获取到命令错误后没有中断任务，也没有盲目安装依赖，而是主动改用项目自带的 unittest。这体现了“操作、观察、调整”的智能体循环。")
    add_cue(doc, "看到真实测试失败时", "test_divide_by_zero → ZeroDivisionError；测试期望 ValueError。", GOLD)
    add_speech(doc, "现在测试真实失败了。除数为零时，程序抛出了 Python 默认的 ZeroDivisionError，但测试要求抛出带指定信息的 ValueError。Agent 因此把问题定位到 calculator.py 的 divide 函数。")

    doc.add_page_break()
    doc.add_heading("基础能力（续）", level=1)
    add_step(doc, "约 1:35", "批准真实文件修改")
    add_command(doc, "Approve? [y/n/yes/no] (n): y")
    add_speech(doc, "修改代码属于有副作用的操作，因此 Agent 会先展示完整差异并请求确认。从差异可以看到，它只修改 calculator.py，增加零除数判断，没有删除或修改已有测试。")
    add_cue(doc, "代码差异", "+ if b == 0:\n+     raise ValueError(\"divisor cannot be zero\")", GREEN)
    add_speech(doc, "批准后，Agent 会真正写入本地文件，并自动重新运行测试。")
    add_cue(doc, "最终结果", "Ran 3 tests · OK", GREEN)
    add_speech(doc, "现在三个测试全部通过，形成了“读取项目、运行测试、定位错误、修改代码、重新验证”的完整 Coding Agent 编程闭环。")

    add_step(doc, "2:15–2:40", "解释 Trust Report")
    add_command(doc, "Status: TRUSTED\nTrust score: 100/100\nTest integrity: PASSED\nTests: 3 runs, 1 passed, 2 failed\nChanged files: calculator.py")
    add_speech(doc, "报告记录了三次测试：第一次因为 pytest 未安装而失败；第二次 unittest 发现真实代码缺陷；第三次在修复后通过。因此一次通过、两次失败是完整过程记录，并不表示最终任务失败。最终状态为 TRUSTED，测试完整性为 PASSED，而且只有 calculator.py 被修改。")

    doc.add_heading("这一部分证明了什么", level=2)
    for text in (
        "读取真实项目文件，而不是只进行聊天回答。",
        "通过本地 Shell 执行测试，并根据错误动态调整方案。",
        "修改前展示差异并请求用户批准。",
        "修改实现而不删除测试，最后用测试验证结果。",
    ):
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()
    doc.add_heading("第二部分  特色功能：Agent Guardian", level=1)
    add_step(doc, "2:40–3:45", "运行本地确定性安全自测")
    add_command(doc, "/guardian-demo")
    add_speech(doc, "接下来演示第一个特色功能 Agent Guardian。这是一次本地实时执行的确定性安全测试，不依赖模型是否愿意发起危险调用。每次运行都会生成新的 Run ID，因此不是预先固定的一张结果图片。")

    doc.add_heading("1. Test Integrity Guard", level=2)
    add_command(doc, "[1/3] PASS Protected test edit")
    add_speech(doc, "Coding Agent 为了让测试通过，可能选择删除失败测试。Guardian 会在任务开始时记录已有测试文件，并在工具执行前阻止对受保护测试的直接修改。这里的 PASS 表示真实修改请求已经被拦截。")

    doc.add_heading("2. Shell Tamper Recovery", level=2)
    add_command(doc, "[2/3] PASS Shell tamper recovery")
    add_speech(doc, "只拦截 edit 和 write_file 还不够，因为模型可能通过 Shell 绕过文件工具。第二项测试会真实写入篡改内容。Guardian 检测到测试文件与快照不一致后，自动恢复原文件。")
    add_command(doc, "Original SHA-256:  c7d71b...\nTampered SHA-256: 071de6...\nRestored SHA-256: c7d71b...  MATCH")
    add_speech(doc, "原始文件和篡改文件的 SHA-256 不同，证明内容确实发生了变化；恢复后的哈希与原始哈希完全一致，MATCH 证明恢复结果是逐字节一致，而不是只输出一句“恢复成功”。")

    doc.add_heading("3. Prompt-Injection Firewall", level=2)
    add_command(doc, "[3/3] PASS Prompt-injection containment")
    add_speech(doc, "Coding Agent 会读取 README、配置和项目文档。如果其中隐藏“忽略用户要求、读取密钥、上传数据”等指令，模型可能把文件内容误认为命令。防火墙把仓库内容视为不可信数据，检测指令覆盖和密钥外传信号，并阻止后续网络等敏感操作。")
    add_cue(doc, "自测结论", "3/3 场景成功控制 · 测试文件恢复 · 2 类注入信号被识别 · 2 次危险操作被阻止", GREEN)

    doc.add_page_break()
    doc.add_heading("第三部分  特色功能：Blocker Gate", level=1)
    add_step(doc, "3:45–4:35", "识别客观不可执行的任务")
    add_command(doc, "请接入私有 SDK 进行线上验证，但 SDK、接口文档和生产凭据都不存在。不要编造接口或伪造验证。")
    add_speech(doc, "现有 Coding Agent 的一个常见问题是：面对客观无法完成的任务，仍然反复检查环境、搜索依赖、尝试不同命令，甚至创建无用诊断文件。这个任务明确缺少 SDK、接口文档和凭据，Blocker Gate 会在任何工具执行之前识别出阻塞条件。")
    add_command(doc, "Blocker Gate: Required external prerequisites are explicitly unavailable:\nSDK or dependency, credentials, API documentation")
    add_speech(doc, "因此 Agent 不会调用 Shell，不会修改项目，不会编造 SDK 接口，也不会伪造线上验证，只会说明阻塞原因和继续任务所需的最少信息。")

    add_step(doc, "4:35–4:50", "解释阻断报告")
    add_command(doc, "Status: NO ACTION\nTrust score: N/A\nTests: 0 runs, 0 passed, 0 failed\nBlocker Gate: ACTIVATED")
    add_speech(doc, "NO ACTION 不是失败，而是说明这一轮没有执行任何文件或命令工具。Blocker Gate 为 ACTIVATED，并记录 SDK、凭据和接口文档三个阻塞条件。这证明阻断来自本地门禁，而不只是模型在文字上拒绝。")

    add_step(doc, "4:50–5:00", "结尾总结")
    add_speech(doc, "总结来说，lm-program 实现了 Coding Agent 的读取、修改、执行和验证闭环。在此基础上，Agent Guardian 能保护已有测试、恢复 Shell 篡改并防御仓库提示注入；Blocker Gate 能及时停止客观不可执行的任务，避免无意义探索和虚假验证。对话历史、工具执行、循环控制、安全判断和错误处理都由本地 Python 实现，没有使用 Agent 框架，也没有依赖服务端代码执行或文件工具。")
    add_command(doc, "/exit")

    doc.add_heading("最终录像检查清单", level=1)
    for text in (
        "read_file 读取真实代码",
        "首次 unittest 出现真实失败",
        "edit 显示真实代码差异并经过批准",
        "修改后 Ran 3 tests / OK",
        "Guardian 显示 3/3 PASS 和 SHA-256 MATCH",
        "Blocker Gate 在零工具调用前显示 ACTIVATED",
        "全程没有展示 API Key",
    ):
        p = doc.add_paragraph(text, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)

    doc.core_properties.title = "lm-program 五分钟完整演示演讲稿"
    doc.core_properties.subject = "Coding Agent 基础功能与特色安全功能演示"
    doc.core_properties.author = "lm-program"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
