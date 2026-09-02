from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "lm-program Agent原理与设计答辩详解.docx"


def font(run, size=None, bold=False, color=None, name="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(6)

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(22)
    title.font.color.rgb = RGBColor.from_string("0B2545")

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "1F4D78"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.line_spacing = 1.05
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(8)


def shade_paragraph(paragraph, fill="F4F6F9"):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def p(doc, text="", label=None):
    para = doc.add_paragraph()
    if label:
        r = para.add_run(label)
        font(r, bold=True, color="1F4D78")
    if text:
        para.add_run(text)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    para.add_run(text)


def number(doc, text):
    para = doc.add_paragraph(style="List Number")
    para.add_run(text)


def code(doc, text):
    para = doc.add_paragraph(style="Code Block")
    r = para.add_run(text.strip())
    font(r, name="Consolas", east_asia="Consolas", size=9)


def callout(doc, text):
    para = doc.add_paragraph()
    shade_paragraph(para)
    r = para.add_run(text)
    font(r, bold=True, color="0B2545")


def qa(doc, question, short, detail, defend, code_text=None):
    doc.add_heading(question, level=2)
    p(doc, " " + short, "先答一句：")
    p(doc, " " + detail, "展开解释：")
    if code_text:
        p(doc, "可以结合下面这段代码讲：", "结合代码：")
        code(doc, code_text)
    p(doc, " " + defend, "设计辩护：")


def build():
    doc = Document()
    setup(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("lm-program Agent 原理与设计答辩详解")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("面向推免考核：理解 Agent 为什么这样运转，以及如何为设计决策辩护")
    font(r, size=11, color="555555")

    doc.add_heading("一、你先建立一个正确心智模型", level=1)
    callout(doc, "Agent 不是一个神秘的新模型，而是“大模型 + 工具 + 上下文 + 循环 + 安全边界”的工程系统。")
    p(
        doc,
        "普通大模型只会根据输入生成文本。Coding Agent 的关键变化是：它把“能做的事情”抽象成工具，"
        "让模型每轮决定下一步要调用哪个工具；程序负责真实执行工具，再把执行结果放回上下文，模型继续判断下一步。"
    )
    p(
        doc,
        "所以你答辩时不要把重点放在“我调用了某个大模型”，而要放在“我怎样组织模型、工具、上下文、安全审批和终止条件”。"
    )

    doc.add_heading("二、30 秒答辩版", level=1)
    callout(
        doc,
        "我的项目是一个终端版 AI 编程 Agent。它兼容 OpenAI 风格接口，通过 ReAct 工具调用循环，让模型能够在本地项目里读文件、改文件、执行命令和运行测试。系统自己实现了上下文管理、工具注册、参数校验、本地工具执行、安全审批、自动验证、任务计划和会话持久化，没有使用 LangChain、AutoGen 这类 Agent 框架。"
    )

    doc.add_heading("三、整体运行流程", level=1)
    for step in [
        "用户在终端输入需求，例如“帮我修复一个 bug 并运行测试”。",
        "main.py 加载配置，包括 API Key、模型名、工作目录和审批策略。",
        "CLI 创建 Agent，Agent 创建或使用 Session。",
        "Session 组装 LLMClient、ToolRegistry、ContextManager、ApprovalManager、MCPManager 等组件。",
        "Agent.run() 把用户消息加入上下文，然后进入 _agentic_loop()。",
        "每一轮把 messages 和 tool_schemas 发给大模型。",
        "如果模型返回文本，就展示给用户；如果模型返回 tool_calls，就执行工具。",
        "ToolRegistry.invoke() 校验参数、走安全审批、触发 hook、执行工具。",
        "工具结果作为 tool message 写回上下文，模型继续下一轮。",
        "如果修改了源码，系统自动发现测试命令并运行验证。",
        "模型不再请求工具时，输出最终回答并结束。",
    ]:
        number(doc, step)

    doc.add_heading("四、最核心代码怎么讲", level=1)
    doc.add_heading("1. 入口 main.py", level=2)
    code(
        doc,
        """
def main(prompt: str | None, cwd: Path | None, configure: bool) -> None:
    config = load_config(cwd=cwd)

    if configure:
        configure_credentials_interactively(config)
        return

    cli = CLI(config)
    if prompt:
        result = asyncio.run(cli.run_single(prompt))
    else:
        asyncio.run(cli.run_interactive())
""",
    )
    p(doc, "这段代码说明入口层只做三件事：加载配置、处理凭据配置、选择运行模式。它没有把 Agent 逻辑写死在 main 里。", "解释：")
    p(doc, "这样做的好处是职责清晰。CLI 负责输入输出，Agent 负责推理循环，Session 负责组件装配。", "辩护：")

    doc.add_heading("2. Agent.run()", level=2)
    code(
        doc,
        """
async def run(self, message: str):
    await self.session.hook_system.trigger_before_agent(message)
    yield AgentEvent.agent_start(message)
    self.session.context_manager.add_user_message(message)

    async for event in self._agentic_loop(message):
        yield event

    await self.session.hook_system.trigger_after_agent(message, final_response)
    yield AgentEvent.agent_end(final_response)
""",
    )
    p(doc, "run() 是一轮用户请求的外层包装：先触发前置 hook，再把用户消息加入上下文，然后进入核心循环，最后触发后置 hook。", "解释：")
    p(doc, "hook 的设计让系统可以扩展，比如记录日志、做审计、接入外部脚本，而不用改 Agent 主逻辑。", "辩护：")

    doc.add_heading("3. _agentic_loop()", level=2)
    code(
        doc,
        """
for turn_num in range(max_turns):
    self.session.increment_turn()
    tool_schemas = self.session.tool_registry.get_schemas()

    async for event in self.session.client.chat_completion(
        self.session.context_manager.get_messages(),
        tools=tool_schemas if tool_schemas else None,
    ):
        if event.type == StreamEventType.TEXT_DELTA:
            response_text += content
        elif event.type == StreamEventType.TOOL_CALL_COMPLETE:
            tool_calls.append(event.tool_call)

    if not tool_calls:
        return

    for tool_call in tool_calls:
        result = await self.session.tool_registry.invoke(...)
""",
    )
    p(doc, "这是项目最关键的 ReAct 循环。每一轮模型要么输出文本，要么要求调用工具。调用工具后，工具结果会回到上下文，下一轮模型就能基于结果继续推理。", "解释：")
    p(doc, "必须用循环，因为编程任务往往不是一步完成的。比如修 bug 需要先跑测试、读报错、搜代码、改文件、再跑测试。单次模型调用做不到可靠闭环。", "辩护：")

    doc.add_heading("4. ToolRegistry.invoke()", level=2)
    code(
        doc,
        """
tool = self.get(name)
validation_errors = tool.validate_params(params)

await hook_system.trigger_before_tool(name, params)
confirmation = await tool.get_confirmation(invocation)
decision = await approval_manager.check_approval(context)

result = await tool.execute(invocation)
await hook_system.trigger_after_tool(name, params, result)
""",
    )
    p(doc, "工具调用不是模型说执行就执行，而是统一经过注册中心。注册中心负责查找工具、校验参数、安全审批、执行和后置 hook。", "解释：")
    p(doc, "这个设计把风险集中在一个边界层里处理，避免每个工具都散落一套审批逻辑，也方便后续扩展新工具。", "辩护：")

    doc.add_heading("五、老师可能重点追问的问题", level=1)
    qa(
        doc,
        "1. 你为什么选择终端形式，而不是做前端页面？",
        "题目要求的是 coding agent 的核心能力，不要求前端页面；终端更贴近开发者真实使用场景。",
        "coding agent 的主要动作是读写文件、执行命令、运行测试，这些本来就是开发者在终端里完成的事情。终端形式能直接展示 Agent 是否真的会操作项目，而不是只展示一个包装界面。",
        "我不是不会做前端，而是这个题目考核重点是 Agent 的核心逻辑。终端方案减少展示层干扰，更能突出我自己实现了上下文管理、工具调用、本地执行和安全审批。",
    )
    qa(
        doc,
        "2. 你为什么不能直接让大模型输出代码，而要做工具调用？",
        "因为大模型本身不能直接访问本地文件和执行命令。",
        "如果只让模型输出代码，用户还要手动复制、粘贴、运行测试。工具调用把这些动作变成程序可执行的能力，模型只需要决定下一步做什么。",
        "工具调用让系统从“建议型助手”变成“执行型助手”。这是 coding agent 和普通聊天机器人的本质区别。",
    )
    qa(
        doc,
        "3. 你为什么没有用 LangChain、AutoGen 这些框架？",
        "题目明确禁止使用 agent 框架或 SDK，所以核心逻辑必须自己实现。",
        "我的项目自己实现了 Tool、ToolRegistry、Agent Loop、ContextManager、ApprovalManager、PersistenceManager 等关键模块。",
        "这样做也有好处：我能解释每个模块为什么存在，而不是把核心逻辑交给黑盒框架。",
    )
    qa(
        doc,
        "4. 你的 Agent 为什么能多步完成任务？",
        "因为它不是单次请求，而是 _agentic_loop 多轮循环。",
        "每轮模型根据当前上下文决定下一步。如果需要文件内容，就调用 read_file；如果要修改，就调用 edit；如果要验证，就调用 shell。工具结果回到上下文后，模型就知道刚才发生了什么。",
        "多步能力来自循环和上下文，而不是模型天然会操作电脑。",
        """
# agent/agent.py
if not tool_calls:
    return

for tool_call in tool_calls:
    result = await self.session.tool_registry.invoke(...)
    tool_call_results.append(...)
""",
    )
    qa(
        doc,
        "5. 你为什么要有 max_turns？",
        "防止 Agent 无限循环。",
        "如果模型一直调用工具但不给最终答案，或者工具结果一直不能解决问题，循环可能停不下来。max_turns 是系统层面的终止条件。",
        "Agent 系统必须设计终止条件，否则自动化能力越强，失控风险越高。",
    )
    qa(
        doc,
        "6. 为什么简单问候不启用工具？",
        "为了减少不必要的工具 schema 和 token 消耗。",
        "如果用户只是说 hello 或 thanks，没有必要把所有工具都传给模型。项目里 _should_enable_tools 会过滤简单消息。",
        "这是一个成本和体验优化：简单对话直接回答，复杂任务再启用工具。",
        """
# agent/agent.py
simple_messages = {"hi", "hello", "hey", "thanks", "ok", "okay"}
return normalized not in simple_messages
""",
    )
    qa(
        doc,
        "7. 为什么要把工具抽象成 Tool 基类？",
        "为了让所有工具有统一接口。",
        "每个工具都有 name、description、kind、schema 和 execute。这样 ToolRegistry 不需要关心具体工具内部怎么实现，只要按统一协议调用。",
        "抽象 Tool 可以降低扩展成本。以后加新工具，只要继承 Tool 并实现 execute，就能注册给模型使用。",
        """
# tools/base.py
class Tool(abc.ABC):
    name: str = "base_tool"
    description: str = "Base tool"
    kind: ToolKind = ToolKind.READ

    @abc.abstractmethod
    async def execute(self, invocation: ToolInvocation) -> ToolResult:
        raise NotImplementedError
""",
    )
    qa(
        doc,
        "8. 为什么要有 ToolKind？",
        "为了区分工具风险等级和用途。",
        "READ 工具一般没有副作用，WRITE、SHELL、NETWORK、MEMORY 通常会改变文件、执行命令、访问网络或写入状态，所以审批逻辑会参考 kind。",
        "这让安全策略不依赖工具名字硬猜，而是有一个明确的工具分类。",
        """
# tools/base.py
class ToolKind(str, Enum):
    READ = "read"
    WRITE = "write"
    SHELL = "shell"
    NETWORK = "network"
    MEMORY = "memory"
    MCP = "mcp"
""",
    )
    qa(
        doc,
        "9. 为什么写操作需要审批，读操作一般不需要？",
        "因为写操作有副作用，可能改坏文件；读操作通常只是获取信息。",
        "项目里 Tool.is_mutating 会把 WRITE、SHELL、NETWORK、MEMORY 视为有副作用的工具。审批系统只对有风险的操作介入。",
        "这样能平衡效率和安全：读文件不频繁打断用户，危险操作必须受控。",
        """
# tools/base.py
def is_mutating(self, params: dict[str, Any]) -> bool:
    return self.kind in {
        ToolKind.WRITE,
        ToolKind.SHELL,
        ToolKind.NETWORK,
        ToolKind.MEMORY,
    }
""",
    )
    qa(
        doc,
        "10. 为什么 EditTool 要求 old_string 精确匹配？",
        "为了减少误改代码的风险。",
        "模型可能对代码位置理解不准。如果按行号或模糊描述修改，很容易改错地方。old_string 精确匹配要求模型先读文件，再拿到原文片段进行替换。",
        "这个设计是保守的。宁愿因为匹配不到而失败，也不要静默改错文件。",
        """
# tools/builtin/edit_file.py
occurrence_count = old_content.count(params.old_string)

if occurrence_count == 0:
    return self._no_match_error(...)

if occurrence_count > 1 and not params.replace_all:
    return ToolResult.error_result(...)
""",
    )
    qa(
        doc,
        "11. 为什么读文件要限制大小、拒绝二进制、带行号？",
        "为了控制上下文长度，并提高模型理解代码的准确性。",
        "超大文件会消耗大量 token，二进制文件对模型没有直接意义，行号则方便定位代码。",
        "工具不是能用就行，还要适合模型使用。读文件工具的输出格式是为后续推理和编辑服务的。",
        """
# tools/builtin/read_file.py
if file_size > self.MAX_FILE_SIZE:
    return ToolResult.error_result(...)

if is_binary_file(path):
    return ToolResult.error_result(...)

formatted_lines.append(f"{i:6}|{line}")
""",
    )
    qa(
        doc,
        "12. 为什么需要 ApprovalManager？",
        "因为 Agent 可以执行真实命令，必须防止危险操作。",
        "比如删除文件、git push、网络下载脚本、修改配置文件，这些都可能造成不可逆影响。ApprovalManager 会根据正则规则、路径和策略判断是否允许。",
        "安全审批是 coding agent 的底线。没有审批的 Agent 不适合操作真实项目。",
        """
# safety/approval.py
requires_confirmation = (
    context.tool_name in {"web_fetch", "web_search"}
    or bool(context.command and is_high_risk_command(context.command))
    or any(is_configuration_path(path) for path in context.affected_paths)
    or context.is_dangerous
)
""",
    )
    qa(
        doc,
        "13. 为什么上下文要自己管理？",
        "因为模型每次请求都需要完整 messages，项目必须决定放哪些历史、工具结果和系统提示词。",
        "ContextManager 保存 user、assistant、tool 消息，并动态构造 system prompt。它还会估算 token，必要时触发压缩或裁剪旧工具输出。",
        "上下文管理决定 Agent 是否能连续工作。没有上下文，模型每一步都会忘记上一步工具结果。",
        """
# context/manager.py
def get_messages(self) -> list[dict[str, Any]]:
    messages = []
    system_prompt = self._build_system_prompt()
    messages.append({"role": "system", "content": system_prompt})
    for item in self._messages:
        messages.append(item.to_dict())
    return messages
""",
    )
    qa(
        doc,
        "14. 为什么需要上下文压缩？",
        "因为长任务会不断产生消息和工具输出，最终超过模型上下文窗口。",
        "当 estimate_context_tokens 超过 context_window 的 80%，项目会调用 ChatCompactor 压缩历史，把之前的工作总结成一段恢复上下文。",
        "压缩的目标不是省略历史，而是保留关键事实，让 Agent 能继续工作。",
        """
# context/manager.py
def needs_compression(self) -> bool:
    context_limit = self.config.model.context_window
    current_tokens = self.estimate_context_tokens()
    return current_tokens > (context_limit * 0.8)
""",
    )
    qa(
        doc,
        "15. 为什么要自动运行测试？",
        "因为 coding agent 改代码后必须验证结果。",
        "项目在 edit 或 write_file 成功修改源码后，会调用 discover_verification_checks 找到测试、lint、typecheck 或 build 命令，然后通过 shell 工具执行。",
        "这是从“能修改”到“能交付”的关键一步。没有验证，Agent 只能说自己改了，不能证明改对了。",
        """
# agent/agent.py
if result.success and tool_call.name in {"edit", "write_file"} and is_source_file(changed_path):
    for check in discover_verification_checks(self.config.cwd):
        verification = await self.session.tool_registry.invoke("shell", ...)
""",
    )
    qa(
        doc,
        "16. 为什么读操作失败可以重试，写操作不自动重试？",
        "因为读操作通常没有副作用，写操作可能已经部分执行。",
        "项目里判断工具不是 mutating 时，失败才会自动重试一次。写文件、shell、网络请求如果自动重试，可能重复写入、重复提交或重复执行危险操作。",
        "这是对副作用的控制。自动化系统不能为了成功率牺牲安全性。",
        """
# agent/agent.py
tool = self.session.tool_registry.get(tool_call.name)
if not result.success and tool and not tool.is_mutating(tool_call.arguments):
    result = await self.session.tool_registry.invoke(...)
""",
    )
    qa(
        doc,
        "17. 为什么要支持会话保存和恢复？",
        "因为真实编程任务可能很长，用户可能中断后继续。",
        "PersistenceManager 会保存 session_id、messages、turn_count 和 token 用量。恢复时重新创建 Session，再把历史消息加载回 ContextManager。",
        "这让 Agent 不只是一次性脚本，而是能持续协作的开发助手。",
        """
# agent/persistence.py
class SessionSnapshot:
    session_id: str
    created_at: datetime
    updated_at: datetime
    turn_count: int
    messages: list[dict[str, Any]]
    total_usage: TokenUsage
""",
    )
    qa(
        doc,
        "18. 为什么会话保存要用原子写入？",
        "为了避免保存过程中程序崩溃导致文件损坏。",
        "它先写临时文件，flush、fsync 后再 os.replace 到正式文件。这样要么旧文件还在，要么新文件完整写入。",
        "这是数据可靠性的设计。虽然项目不大，但会话数据损坏会影响恢复体验。",
        """
# agent/persistence.py
fd, temporary_path = tempfile.mkstemp(...)
json.dump(snapshot.to_dict(), fp, indent=2)
fp.flush()
os.fsync(fp.fileno())
os.replace(temporary_path, file_path)
""",
    )
    qa(
        doc,
        "19. MCP 为什么要做？",
        "为了让 Agent 能接入外部工具能力。",
        "MCPManager 读取配置，连接 MCP server，然后把外部工具包装成 MCPTool 注册到 ToolRegistry。这样模型能像调用内置工具一样调用外部工具。",
        "这体现了扩展性。内置工具解决基础能力，MCP 让系统可以继续扩展。",
        """
# tools/mcp/mcp_manager.py
for tool_info in client.tools:
    mcp_tool = MCPTool(
        tool_info=tool_info,
        client=client,
        name=f"{client.name}__{tool_info.name}",
    )
    registry.register_mcp_tool(mcp_tool)
""",
    )
    qa(
        doc,
        "20. 这个项目有什么不足？你怎么优化？",
        "可以从拼写 bug、类型一致性、验证配置、安全规则和用户体验回答。",
        "目前 ApprovalPolicy.AUTO_EDIT 写成 auto-edut，ToolCall.arguments 的类型标注和实际 dict 使用不完全一致，自动验证规则比较简单，审批主要依赖正则。",
        "我会先修明确 bug 并补测试，再增加项目级验证命令配置；安全规则可以从简单正则逐步升级成命令解析；CLI 中文乱码也可以统一处理。",
        """
# config/config.py
class ApprovalPolicy(str, Enum):
    AUTO_EDIT = "auto-edut"  # 疑似应为 auto-edit

# client/response.py
class ToolCall:
    arguments: str = ""  # 实际 parse 后通常按 dict 使用
""",
    )

    doc.add_heading("六、如果老师让你手撕代码", level=1)
    doc.add_heading("1. 手撕 Agent Loop", level=2)
    code(
        doc,
        """
async def agent_loop(user_message, llm, tool_registry, tool_schemas, max_turns=20):
    messages = [{"role": "user", "content": user_message}]

    for _ in range(max_turns):
        response = await llm.chat(messages, tools=tool_schemas)

        messages.append({
            "role": "assistant",
            "content": response.text,
            "tool_calls": response.tool_calls,
        })

        if not response.tool_calls:
            return response.text

        for call in response.tool_calls:
            result = await tool_registry.invoke(call.name, call.arguments)
            messages.append({
                "role": "tool",
                "tool_call_id": call.id,
                "content": result,
            })

    return "任务未在最大轮数内完成"
""",
    )
    p(doc, "讲解时说：这段代码体现了 Agent 的本质，即模型根据上下文决定行动，工具执行行动，结果回到上下文，循环推进。", "怎么讲：")

    doc.add_heading("2. 手撕 ToolRegistry", level=2)
    code(
        doc,
        """
class Tool:
    name = ""

    async def execute(self, params):
        raise NotImplementedError


class ToolRegistry:
    def __init__(self):
        self.tools = {}

    def register(self, tool):
        self.tools[tool.name] = tool

    async def invoke(self, name, params):
        tool = self.tools.get(name)
        if not tool:
            return {"ok": False, "error": "unknown tool"}
        return await tool.execute(params)
""",
    )
    p(doc, "讲解时说：真实项目在这个基础上增加了 Pydantic 参数校验、安全审批、hook、MCP 工具和异常处理。", "怎么讲：")

    doc.add_heading("七、答辩时的万能组织方式", level=1)
    callout(doc, "老师问任何模块，都按“它解决什么问题 -> 代码在哪里 -> 为什么这样设计 -> 有什么风险和优化”四步回答。")
    bullet(doc, "它解决什么问题：比如 ToolRegistry 解决模型工具调用如何落地执行的问题。")
    bullet(doc, "代码在哪里：指出具体文件，例如 tools/registry.py。")
    bullet(doc, "为什么这样设计：统一管理、方便扩展、安全边界清晰。")
    bullet(doc, "风险和优化：比如参数类型、异常处理、审批规则、测试覆盖。")

    doc.add_heading("八、最后可以背的一段", level=1)
    callout(
        doc,
        "我这个项目的核心是自己实现一个 coding agent 的运行闭环。模型不直接操作系统，而是通过 tool calling 表达意图；ToolRegistry 负责把意图变成本地执行；ApprovalManager 控制危险操作；ContextManager 保存多轮消息和工具结果；Agent Loop 让模型能根据观察结果继续行动；自动验证让代码修改后能跑测试确认结果。这个设计的重点是把大模型的不确定推理能力，放进一个有边界、有状态、可验证的工程系统里。"
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("lm-program Agent 原理与设计答辩详解")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
