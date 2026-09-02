from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "lm-program项目面试考核讲解-老师追问版.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])


def style_run(run, bold=False, color=None, size=None, font="Calibri") -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_code(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = "Code Block"
    run = paragraph.add_run(code.strip())
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)


def add_kv_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table, [1700, 7660])
    set_cell_margins(table)
    hdr = table.rows[0].cells
    hdr[0].text = "模块"
    hdr[1].text = "考核时怎么讲"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                style_run(r, bold=True, color="0B2545")
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)


def add_qa_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table, [3100, 6260])
    set_cell_margins(table)
    hdr = table.rows[0].cells
    hdr[0].text = "可能问题"
    hdr[1].text = "参考回答"
    for cell in hdr:
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                style_run(r, bold=True, color="0B2545")
    for q, a in rows:
        cells = table.add_row().cells
        cells[0].text = q
        cells[1].text = a


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(22)
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    code_style = styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.1


def build_document() -> None:
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("lm-program 项目面试考核讲解稿")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("终端版 AI 编程 Agent：项目介绍、核心架构、代码手撕与面试问答")
    style_run(run, color="555555", size=11)

    doc.add_heading("一、30 秒项目介绍", level=1)
    doc.add_paragraph(
        "这个项目是一个基于 Python 的终端 AI 编程 Agent。用户在命令行输入自然语言需求后，"
        "系统会调用 OpenAI 兼容的大模型进行推理，并通过工具系统完成读文件、改文件、执行命令、"
        "保存会话、任务规划、MCP 扩展和安全审批等操作。它的核心不是单纯聊天，而是让模型真正参与本地代码开发。"
    )

    doc.add_heading("二、核心流程", level=1)
    add_numbered(
        doc,
        [
            "用户通过 CLI 输入 prompt 或进入交互模式。",
            "main.py 加载配置，创建 CLI 和 Agent。",
            "Agent 将用户消息写入 ContextManager。",
            "LLMClient 把上下文和工具 schema 发给大模型。",
            "模型返回文本，或者返回 tool_calls。",
            "ToolRegistry 校验参数、安全审批并执行工具。",
            "工具结果回写上下文，模型继续推理。",
            "如果源码被修改，自动发现并执行项目测试。",
            "最终输出回答，并支持保存或恢复会话。",
        ],
    )

    doc.add_heading("三、模块职责速查", level=1)
    add_kv_table(
        doc,
        [
            ("main.py", "命令行入口。负责配置、单次运行、交互运行，以及 /help、/tools、/save、/resume 等命令。"),
            ("agent/agent.py", "Agent 主循环。负责调用大模型、收集工具调用、执行工具、自动验证、输出事件。"),
            ("agent/session.py", "会话运行时容器。组装 LLMClient、工具注册中心、上下文、MCP、审批、hook、技能系统。"),
            ("client/llm_client.py", "OpenAI 兼容客户端。负责流式调用、解析文本增量、解析工具调用和 token 用量。"),
            ("tools/registry.py", "工具注册中心。注册工具、生成 schema、校验参数、接入审批、执行工具。"),
            ("tools/base.py", "工具抽象层。定义 Tool、ToolResult、ToolConfirmation、FileDiff 等基础模型。"),
            ("context/manager.py", "上下文管理。保存 user/assistant/tool 消息，构造系统提示词，必要时裁剪工具输出。"),
            ("safety/approval.py", "安全审批。根据命令风险、路径、配置文件和审批策略决定是否执行。"),
            ("context/verification.py", "自动验证。根据项目类型发现测试、lint、typecheck、build 等命令。"),
            ("agent/persistence.py", "会话持久化。把消息、用量、时间和 session_id 原子写入本地 JSON。"),
        ],
    )

    doc.add_heading("四、核心代码讲解", level=1)
    doc.add_heading("1. CLI 入口", level=2)
    doc.add_paragraph("位置：main.py。它决定是配置模式、单次 prompt 模式，还是交互式模式。")
    add_code(
        doc,
        """
@click.command()
@click.argument("prompt", required=False)
@click.option("--cwd", "-c", ...)
@click.option("--configure", is_flag=True, ...)
def main(prompt, cwd, configure):
    config = load_config(cwd=cwd)

    if configure:
        configure_credentials_interactively(config)
        return

    cli = CLI(config)
    if prompt:
        asyncio.run(cli.run_single(prompt))
    else:
        asyncio.run(cli.run_interactive())
""",
    )
    doc.add_paragraph("考核说法：CLI 只负责交互和命令分发，真正的智能体逻辑交给 Agent。")

    doc.add_heading("2. Agent 主循环", level=2)
    doc.add_paragraph("位置：agent/agent.py。最重要的是 _agentic_loop，它实现 ReAct 风格循环。")
    add_code(
        doc,
        """
async for event in self.session.client.chat_completion(
    self.session.context_manager.get_messages(),
    tools=tool_schemas if tool_schemas else None,
):
    if event.type == StreamEventType.TEXT_DELTA:
        response_text += event.text_delta.content
    elif event.type == StreamEventType.TOOL_CALL_COMPLETE:
        tool_calls.append(event.tool_call)
""",
    )
    doc.add_paragraph("这里的关键点是：模型每轮可以选择直接回答，也可以选择调用工具。")

    add_code(
        doc,
        """
result = await self.session.tool_registry.invoke(
    tool_call.name,
    tool_call.arguments,
    self.config.cwd,
    self.session.hook_system,
    self.session.approval_manager,
)
""",
    )
    doc.add_paragraph("考核说法：Agent 不直接读写文件，而是把工具调用交给 ToolRegistry，统一校验和审批。")

    doc.add_heading("3. 工具注册中心", level=2)
    doc.add_paragraph("位置：tools/registry.py。它是模型能力和真实系统能力之间的边界。")
    add_code(
        doc,
        """
validation_errors = tool.validate_params(params)
if validation_errors:
    return ToolResult.error_result(...)

confirmation = await tool.get_confirmation(invocation)
decision = await approval_manager.check_approval(context)

result = await tool.execute(invocation)
""",
    )
    doc.add_paragraph("考核说法：工具执行前会先校验参数，再走审批策略，最后才真正执行。")

    doc.add_heading("4. 读文件工具", level=2)
    doc.add_paragraph("位置：tools/builtin/read_file.py。它体现了 Agent 工具设计的边界意识。")
    add_bullets(
        doc,
        [
            "先检查路径是否存在、是否是文件。",
            "限制最大文件大小，避免读取超大文件。",
            "拒绝二进制文件，避免乱码和 token 浪费。",
            "支持 offset 和 limit 分段读取。",
            "输出带行号，方便模型定位和修改。",
        ],
    )
    add_code(
        doc,
        """
if not path.exists():
    return ToolResult.error_result(f"File not found: {path}")

if is_binary_file(path):
    return ToolResult.error_result("Cannot read binary file")

for i, line in enumerate(selected_lines, start=start_idx + 1):
    formatted_lines.append(f"{i:6}|{line}")
""",
    )

    doc.add_heading("5. 编辑文件工具", level=2)
    doc.add_paragraph("位置：tools/builtin/edit_file.py。它采用 old_string 到 new_string 的精确替换。")
    add_code(
        doc,
        """
occurrence_count = old_content.count(params.old_string)

if occurrence_count == 0:
    return self._no_match_error(...)

if occurrence_count > 1 and not params.replace_all:
    return ToolResult.error_result(...)
""",
    )
    doc.add_paragraph(
        "考核说法：这种设计可以减少 AI 改代码时的误伤。匹配不到会提示相似行，匹配多次会要求提供更精确上下文。"
    )

    doc.add_heading("6. 安全审批", level=2)
    doc.add_paragraph("位置：safety/approval.py。它按照命令风险决定是否自动执行、要求确认或拒绝。")
    add_code(
        doc,
        """
if is_dangerous_command(command):
    return ApprovalDecision.REJECTED

if context.command and is_high_risk_command(context.command):
    return ApprovalDecision.NEEDS_CONFIRMATION
""",
    )
    doc.add_paragraph("考核说法：因为 Agent 能执行真实命令，所以必须有安全边界。")

    doc.add_heading("7. 自动验证", level=2)
    doc.add_paragraph("位置：agent/agent.py 和 context/verification.py。源码修改成功后自动运行测试。")
    add_code(
        doc,
        """
if result.success and tool_call.name in {"edit", "write_file"} and is_source_file(changed_path):
    for check in discover_verification_checks(self.config.cwd):
        verification = await self.session.tool_registry.invoke("shell", ...)
""",
    )
    doc.add_paragraph("本项目当前测试命令是：uv run python -m unittest discover -s tests。实际运行结果：25 个测试全部通过。")

    doc.add_heading("五、手撕代码：简化版 Agent Loop", level=1)
    add_code(
        doc,
        """
async def agent_loop(user_message):
    messages.append({"role": "user", "content": user_message})

    for _ in range(max_turns):
        response = await llm.chat(messages, tools=tool_schemas)

        if response.text:
            print(response.text)

        if not response.tool_calls:
            return response.text

        messages.append({
            "role": "assistant",
            "content": response.text,
            "tool_calls": response.tool_calls,
        })

        for call in response.tool_calls:
            result = await tool_registry.invoke(call.name, call.arguments)
            messages.append({
                "role": "tool",
                "tool_call_id": call.id,
                "content": result,
            })
""",
    )
    doc.add_paragraph("讲解重点：这段代码体现的是“模型决策，程序执行，结果反馈，再继续推理”的循环。")

    doc.add_heading("六、手撕代码：简化版工具注册中心", level=1)
    add_code(
        doc,
        """
class Tool:
    name = ""
    description = ""

    async def execute(self, params):
        raise NotImplementedError


class ToolRegistry:
    def __init__(self):
        self.tools = {}

    def register(self, tool):
        self.tools[tool.name] = tool

    async def invoke(self, name, params):
        if name not in self.tools:
            return {"ok": False, "error": "unknown tool"}

        tool = self.tools[name]
        return await tool.execute(params)
""",
    )
    doc.add_paragraph("可以继续补一个 ReadFileTool：")
    add_code(
        doc,
        """
class ReadFileTool(Tool):
    name = "read_file"
    description = "Read a text file"

    async def execute(self, params):
        path = params["path"]
        with open(path, "r", encoding="utf-8") as f:
            return {"ok": True, "content": f.read()}
""",
    )

    doc.add_heading("七、面试高频问答", level=1)
    add_qa_table(
        doc,
        [
            ("这个项目解决什么问题？", "让大模型不只是聊天，而是能操作本地代码项目，完成读代码、改代码、运行测试和保存进度。"),
            ("这个项目和普通 ChatGPT 调 API 有什么区别？", "普通聊天只返回文本；这个项目把文件读写、shell、搜索、任务管理等能力封装成工具，让模型可以通过 tool_calls 操作真实项目。"),
            ("项目的入口在哪里？", "入口是 main.py 里的 main()。它负责加载配置、检查 API Key，然后根据是否传入 prompt 决定单次运行或交互运行。"),
            ("为什么 main.py 里要用 asyncio.run？", "因为 Agent、LLMClient、工具调用和 MCP 连接都是异步设计，asyncio.run 用来启动异步事件循环。"),
            ("单次运行和交互运行有什么区别？", "单次运行传入一个 prompt 后执行完就退出；交互运行会进入 while 循环，支持连续对话和 /tools、/save、/resume 等命令。"),
            ("为什么需要 Agent Loop？", "因为一次模型调用往往不能完成复杂任务，需要多轮推理、工具调用和结果反馈。"),
            ("Agent.run() 做了什么？", "它先触发 before_agent hook，把用户消息加入上下文，然后进入 _agentic_loop，最后触发 after_agent hook 并结束本轮。"),
            ("_agentic_loop 的核心逻辑是什么？", "每轮先检查上下文是否需要压缩，再把 messages 和 tool_schemas 发给模型；如果模型返回工具调用，就执行工具并把结果写回上下文，继续下一轮。"),
            ("为什么要限制 max_turns？", "防止模型陷入无限循环，比如一直调用工具但不给最终答案。max_turns 是一个安全兜底。"),
            ("什么情况下不会启用工具？", "Agent 里 _should_enable_tools 会判断简单问候，例如 hi、hello、thanks，这类输入不需要工具，避免浪费 token 和无意义调用。"),
            ("为什么需要工具注册中心？", "统一管理模型可调用能力，把工具转成 schema，并负责参数校验、安全审批和执行。"),
            ("工具 schema 是干什么的？", "schema 会告诉大模型工具名字、描述和参数结构。模型根据 schema 生成符合格式的 tool_calls。"),
            ("工具参数怎么校验？", "每个工具用 Pydantic BaseModel 定义参数，Tool.validate_params 会实例化模型，如果字段缺失或类型不对就返回错误。"),
            ("ToolKind 有什么作用？", "ToolKind 标记工具类型，比如 READ、WRITE、SHELL、NETWORK、MEMORY、MCP。审批系统会根据类型判断工具是否有副作用。"),
            ("读文件工具为什么要加 offset 和 limit？", "为了支持大文件分段读取，避免一次把超长文件塞进上下文，也方便模型按行定位。"),
            ("为什么读文件输出要带行号？", "大模型修改代码时需要精确定位，带行号能帮助它引用和理解代码位置。"),
            ("为什么不能读二进制文件？", "图片、可执行文件等二进制内容对模型不可读，还会造成乱码和 token 浪费，所以工具会拒绝。"),
            ("EditTool 为什么要求 old_string 精确匹配？", "这是为了做小范围、安全的代码修改。只有精确匹配到目标文本，才替换成 new_string，避免误改无关代码。"),
            ("如果 old_string 匹配到多处怎么办？", "如果 replace_all=False，就返回错误，要求提供更具体的上下文；如果明确 replace_all=True，才会替换全部。"),
            ("FileDiff 的作用是什么？", "FileDiff 用来生成修改前后的 unified diff，方便展示给用户审批，也方便理解文件变化。"),
            ("安全怎么保证？", "ApprovalManager 会识别危险命令、高风险命令、配置文件修改和网络访问，根据审批策略决定是否执行。"),
            ("审批策略有哪些？", "主要有 on-request、on-failure、auto、auto-edit、never、yolo。不同策略决定安全命令、高风险命令是否自动执行或需要确认。"),
            ("哪些操作会被认为高风险？", "删除文件、git commit、git push、git config、curl/wget、pip/npm install、网页搜索或抓取、修改 .env/config 等配置文件。"),
            ("为什么 web_search 也要审批？", "因为它会访问网络，可能暴露查询内容，也可能引入不可信外部信息，所以默认属于需要确认的网络行为。"),
            ("项目如何防止危险命令？", "safety/approval.py 中有 DANGEROUS_PATTERNS，例如 rm -rf /、shutdown、curl | bash 等，匹配到会拒绝或要求确认。"),
            ("上下文太长怎么办？", "ContextManager 会估算 token，当接近上下文窗口上限时触发压缩，并裁剪旧工具输出。"),
            ("ContextManager 里保存哪些消息？", "保存 user、assistant 和 tool 三类消息。发送给模型时还会动态构造 system prompt。"),
            ("system prompt 怎么生成？", "ContextManager 调用 prompts.system.get_system_prompt，把配置、工具、技能、用户记忆和项目概览拼成系统提示词。"),
            ("项目上下文是怎么来的？", "context/project.py 会扫描根目录、语言、依赖文件、README 摘要等，形成一个 bounded project overview，注入系统提示词。"),
            ("为什么要裁剪旧工具输出？", "工具输出可能很长，比如读大文件或测试日志。旧输出保留太多会占满上下文，所以会在安全范围内替换成占位文本。"),
            ("怎么证明改动正确？", "源码被 edit 或 write_file 修改后，会自动发现测试命令并运行，把失败结果反馈给模型继续修。"),
            ("自动验证是怎么发现测试命令的？", "context/verification.py 会检查 pyproject.toml、package.json、tests 目录等。Python 项目有 tests 目录时会运行 unittest discover。"),
            ("为什么只在改源码后自动验证？", "因为只有源码修改才最需要质量闭环。普通读文件或查询不会改变代码，没有必要每次都跑测试。"),
            ("如果测试失败怎么办？", "测试输出会拼接到工具结果中返回给模型，模型可以根据失败信息继续修改，形成修复闭环。"),
            ("MCP 的作用是什么？", "MCP 允许项目接入外部工具服务器，把外部能力注册成模型可调用工具。"),
            ("MCP 工具怎么注册进系统？", "Session.initialize 会初始化 MCPManager，连接配置里的 MCP server，然后 MCPManager.register_tools 把外部工具包装成 MCPTool 注册到 ToolRegistry。"),
            ("Hook 系统有什么作用？", "HookSystem 可以在 before_agent、after_agent、before_tool、after_tool、on_error 等时机执行外部命令或脚本，用来扩展流程。"),
            ("会话怎么恢复？", "PersistenceManager 把 session_id、消息、时间和 token 用量写入本地 JSON，支持 /save、/resume、/checkpoint。"),
            ("为什么保存会话要用原子写入？", "persistence.py 先写临时文件，再 os.replace 替换正式文件，避免程序中途崩溃导致 JSON 写坏。"),
            ("任务计划系统有什么用？", ".tasks/ 目录会保存任务、状态、todo、重试次数和依赖关系，适合多步骤开发任务中断后继续执行。"),
            ("Task 的状态有哪些？", "pending、in_progress、paused、failed、completed。这样可以区分未开始、进行中、暂停、失败和已完成。"),
            ("为什么读操作失败可以自动重试，写操作不自动重试？", "读操作通常没有副作用，重试安全；写文件、shell、网络等可能已经部分执行，自动重试可能造成重复修改或风险。"),
            ("LLMClient 为什么把 SDK 的 max_retries 设为 0？", "项目自己实现了重试逻辑。如果 SDK 也重试，会把一次请求放大成多次等待，导致超时更难控制。"),
            ("流式响应是怎么处理的？", "LLMClient 把模型返回拆成 StreamEvent，比如 TEXT_DELTA、TOOL_CALL_DELTA、TOOL_CALL_COMPLETE，CLI 再边收边展示。"),
            ("parse_tool_call_arguments 有什么兜底？", "如果工具参数 JSON 解析失败，就返回 raw_arguments，避免整个流程直接崩掉，方便把错误反馈给模型。"),
            ("项目用了哪些主要依赖？", "click 做 CLI，rich 做终端展示，openai/httpx 调模型，pydantic 做参数校验，platformdirs 管理配置路径，fastmcp 支持 MCP。"),
            ("配置从哪里加载？", "先加载用户级 config.toml，再合并项目级 .ai-agent/config.toml，还支持 .env、API_KEY 和 BASE_URL 环境变量。"),
            ("AGENT.MD 是什么？", "项目根目录如果有 AGENT.MD，会作为 developer_instructions 加载，用来给 Agent 注入项目级开发规范。"),
            ("这个项目当前有什么明显 bug？", "ApprovalPolicy.AUTO_EDIT 的值写成了 auto-edut，应该是 auto-edit。CLI 某些输出里也有乱码字符，可能是编码问题。"),
            ("如果让你优化这个项目，你会怎么做？", "先修配置拼写和类型标注，再加强文件写入原子性、完善测试覆盖、支持用户自定义验证命令，并优化 CLI 中文输出。"),
            ("如果老师让你现场手撕，你写哪块？", "优先写简化版 Agent Loop 或 ToolRegistry，因为这两个最能体现项目核心思想：模型决策、工具执行、结果反馈。"),
            ("这个项目的设计模式像什么？", "整体像插件化架构加事件驱动。工具是可插拔组件，Agent 通过事件把文本、工具调用、错误和验证过程交给 TUI 展示。"),
            ("为什么说它是 ReAct 风格？", "ReAct 指 Reasoning + Acting。模型先推理，再选择 action，也就是工具调用；工具返回 observation 后，模型继续推理。"),
            ("项目有哪些边界或不足？", "它依赖大模型能力，工具调用准确性不能百分百保证；审批规则主要靠正则；自动验证规则较简单；MCP 连接失败时能力会减少。"),
            ("怎么介绍你对代码的理解？", "可以说我按入口 main.py、运行时 Session、核心 Agent Loop、工具注册中心、安全审批、上下文管理、自动验证这条主线读代码。"),
        ],
    )

    doc.add_heading("八、可以主动提的优化点", level=1)
    add_bullets(
        doc,
        [
            "config/config.py 中 ApprovalPolicy.AUTO_EDIT 的值写成了 auto-edut，疑似拼写错误，应改成 auto-edit 并补测试。",
            "ToolCall 类型中 arguments 标注为 str，但实际 parse_tool_call_arguments 返回 dict[str, Any]，类型可以统一。",
            "EditTool 写文件前可以增加更强的编码兼容和原子写入，避免中途失败损坏文件。",
            "CLI 输出里部分中文符号显示异常，可能是编码或拷贝问题，可以统一修复。",
            "自动验证目前依赖简单规则，可以扩展为读取项目配置中的自定义验证命令。",
        ],
    )

    doc.add_heading("九、最后总结", level=1)
    doc.add_paragraph(
        "这个项目的核心价值在于把大模型接入真实开发环境。它通过 Agent Loop 实现多轮推理，"
        "通过 ToolRegistry 管理工具能力，通过 ApprovalManager 控制风险，通过 ContextManager 管理长上下文，"
        "通过 PersistenceManager 支持会话恢复，并通过自动验证形成修改代码后的质量闭环。"
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("lm-program 面试考核讲解稿")

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(OUT)
