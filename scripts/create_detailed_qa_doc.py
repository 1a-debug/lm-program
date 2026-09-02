from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "lm-program项目面试考核逐题详解版.docx"


def set_font(run, name="Calibri", east_asia="Microsoft YaHei", size=None, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title.font.size = Pt(22)
    title.font.color.rgb = RGBColor.from_string("0B2545")

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    code_style = doc.styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code_style.font.size = Pt(9)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.05


def para(doc, text="", bold_label=None):
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label)
        set_font(r, bold=True, color="1F4D78")
    if text:
        p.add_run(text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def code(doc, text):
    p = doc.add_paragraph(style="Code Block")
    r = p.add_run(text.strip())
    set_font(r, name="Consolas", east_asia="Consolas", size=9)


QUESTIONS = [
    {
        "q": "1. 这个项目到底是干什么的？",
        "idea": "它是一个终端 AI 编程 Agent。重点不是聊天，而是让大模型通过工具操作本地代码项目。",
        "code": """
# pyproject.toml
[project]
name = "lm-program"
description = "A terminal-based OpenAI-compatible coding agent with tools, approvals, MCP, and session persistence."
""",
        "explain": "从项目描述就能看出几个关键词：terminal-based 表示终端运行，OpenAI-compatible 表示兼容 OpenAI 风格接口，coding agent 表示面向代码开发，tools、approvals、MCP、session persistence 分别对应工具、安全审批、外部扩展和会话保存。",
        "oral": "我会把它理解成一个简化版 Codex。用户输入自然语言，它调用大模型思考，再用工具读代码、改代码、跑命令，最后把结果反馈给用户。",
    },
    {
        "q": "2. 项目启动以后第一步做什么？",
        "idea": "先从 main.py 进入，加载配置，检查 API Key，然后决定是配置模式、单次执行还是交互模式。",
        "code": """
# main.py
def main(prompt: str | None, cwd: Path | None, configure: bool) -> None:
    config = load_config(cwd=cwd)

    if configure:
        configure_credentials_interactively(config)
        return

    if not config.api_key:
        ...

    cli = CLI(config)
    if prompt:
        result = asyncio.run(cli.run_single(prompt))
    else:
        asyncio.run(cli.run_interactive())
""",
        "explain": "这里 main() 是命令行入口。load_config 负责加载用户配置和项目配置。configure 为 true 时进入 API Key 配置流程。如果用户直接传 prompt，就调用 run_single；否则调用 run_interactive 进入连续对话。",
        "oral": "入口层不做复杂业务，只负责启动流程和模式选择。真正的 Agent 能力在后面的 Agent 和 Session 里。",
    },
    {
        "q": "3. 为什么项目里大量使用 async/await？",
        "idea": "因为大模型流式响应、MCP 连接、工具执行都可能是 I/O 操作，用异步可以避免阻塞。",
        "code": """
# main.py
if prompt:
    result = asyncio.run(cli.run_single(prompt))
else:
    asyncio.run(cli.run_interactive())

# agent/agent.py
async def run(self, message: str):
    async for event in self._agentic_loop(message):
        yield event
""",
        "explain": "asyncio.run 用来启动异步事件循环。Agent.run 是异步生成器，可以一边接收模型输出，一边向 TUI 发送事件。这样用户能看到流式输出，而不是等全部完成后一次性显示。",
        "oral": "这个项目的请求模型、流式输出和工具调用都是异步场景，所以用 async/await 更合适。",
    },
    {
        "q": "4. Session 在项目里起什么作用？",
        "idea": "Session 是一次会话的运行时容器，负责把模型客户端、工具、上下文、安全审批、MCP 等模块装配起来。",
        "code": """
# agent/session.py
class Session:
    def __init__(self, config: Config):
        self.client = LLMClient(config=config)
        self.tool_registry = create_default_registry(config)
        self.skill_registry = SkillRegistry(config)
        self.mcp_manager = MCPManager(self.config)
        self.chat_compactor = ChatCompactor(self.client)
        self.approval_manager = ApprovalManager(...)
        self.hook_system = HookSystem(config)
""",
        "explain": "Agent 不直接 new 所有组件，而是通过 Session 拿到运行时能力。这样结构更清晰：Agent 负责流程控制，Session 负责资源组织。",
        "oral": "我理解 Session 就像一次对话的总控台，它里面有模型、有工具、有上下文、有安全策略。",
    },
    {
        "q": "5. Agent Loop 是什么？为什么它是核心？",
        "idea": "Agent Loop 是多轮推理和工具调用循环。模型先判断要不要调用工具，工具执行后结果再返回模型，直到最终回答。",
        "code": """
# agent/agent.py
for turn_num in range(max_turns):
    response_text = ""
    tool_calls = []

    async for event in self.session.client.chat_completion(...):
        if event.type == StreamEventType.TEXT_DELTA:
            response_text += content
        elif event.type == StreamEventType.TOOL_CALL_COMPLETE:
            tool_calls.append(event.tool_call)

    if not tool_calls:
        return

    for tool_call in tool_calls:
        result = await self.session.tool_registry.invoke(...)
""",
        "explain": "普通聊天只调用一次模型。Agent 不一样，它要根据模型的工具调用结果继续思考，所以必须有循环。max_turns 是兜底，防止模型一直调用工具停不下来。",
        "oral": "这段是项目最核心的地方，可以概括成：模型决策、工具执行、结果回传、继续推理。",
    },
    {
        "q": "6. 大模型怎么知道有哪些工具可以用？",
        "idea": "ToolRegistry 会把工具转换成 OpenAI function calling schema，传给 LLMClient。",
        "code": """
# tools/registry.py
def get_schemas(self) -> list[dict[str, Any]]:
    return [tool.to_openai_schema() for tool in self.get_tools()]

# agent/agent.py
tool_schemas = self.session.tool_registry.get_schemas()
await self.session.client.chat_completion(messages, tools=tool_schemas)
""",
        "explain": "每个工具都有 name、description 和 parameters。模型看到 schema 后，就能生成对应 tool_call，例如 read_file(path='main.py')。",
        "oral": "工具 schema 相当于给模型一份菜单，告诉它能点哪些工具、每个工具要传什么参数。",
    },
    {
        "q": "7. 工具调用真正是在哪里执行的？",
        "idea": "真正执行在 ToolRegistry.invoke。它先找工具，再校验参数，再安全审批，最后执行 execute。",
        "code": """
# tools/registry.py
tool = self.get(name)
validation_errors = tool.validate_params(params)
await hook_system.trigger_before_tool(name, params)

confirmation = await tool.get_confirmation(invocation)
decision = await approval_manager.check_approval(context)

result = await tool.execute(invocation)
await hook_system.trigger_after_tool(name, params, result)
""",
        "explain": "这段代码体现了工具执行链路。它不是模型说执行就立刻执行，而是经过 registry 统一管理。这样可以集中做校验、审批、hook 和错误处理。",
        "oral": "ToolRegistry 是模型和真实系统之间的安全边界。模型只发意图，程序决定能不能执行、怎么执行。",
    },
    {
        "q": "8. 为什么每个工具要用 Pydantic 定义参数？",
        "idea": "Pydantic 可以自动校验参数类型、必填字段和约束，避免模型传错参数导致工具崩溃。",
        "code": """
# tools/builtin/read_file.py
class ReadFileParams(BaseModel):
    path: str = Field(...)
    offset: int = Field(1, ge=1)
    limit: int | None = Field(None, ge=1)

# tools/base.py
def validate_params(self, params):
    schema(**params)
""",
        "explain": "比如 offset 必须大于等于 1，limit 如果传了也必须大于等于 1。模型生成 JSON 参数可能不稳定，所以工具执行前必须校验。",
        "oral": "Pydantic 在这里相当于工具参数的门卫，参数不合法就不给执行。",
    },
    {
        "q": "9. ReadFileTool 为什么不直接 read_text 完事？",
        "idea": "因为 Agent 读文件要考虑安全、上下文长度和可读性，所以加了文件大小限制、二进制判断、分页和行号。",
        "code": """
# tools/builtin/read_file.py
if file_size > self.MAX_FILE_SIZE:
    return ToolResult.error_result(...)

if is_binary_file(path):
    return ToolResult.error_result(...)

selected_lines = lines[start_idx:end_idx]
for i, line in enumerate(selected_lines, start=start_idx + 1):
    formatted_lines.append(f"{i:6}|{line}")
""",
        "explain": "大模型上下文有限，不能随便塞超大文件。二进制文件也没有直接阅读价值。行号能帮助模型后续定位修改位置。",
        "oral": "这个工具体现了工程细节：不仅能读，还要读得安全、可控、方便模型理解。",
    },
    {
        "q": "10. EditTool 为什么使用 old_string/new_string 替换？",
        "idea": "这是为了让修改更精确，避免模型按行号或模糊描述误改代码。",
        "code": """
# tools/builtin/edit_file.py
occurrence_count = old_content.count(params.old_string)

if occurrence_count == 0:
    return self._no_match_error(...)

if occurrence_count > 1 and not params.replace_all:
    return ToolResult.error_result(...)

new_content = old_content.replace(params.old_string, params.new_string, 1)
""",
        "explain": "old_string 必须完全匹配，包括空格和缩进。找不到就不改；匹配多处也不默认改。这样可以降低 AI 修改代码时的误伤风险。",
        "oral": "它不是让模型随便重写文件，而是要求模型给出精确上下文，属于比较安全的编辑方式。",
    },
    {
        "q": "11. 安全审批是怎么工作的？",
        "idea": "ApprovalManager 根据工具类型、命令风险、路径和审批策略判断是否允许执行。",
        "code": """
# safety/approval.py
requires_confirmation = (
    context.tool_name in {"web_fetch", "web_search"}
    or bool(context.command and is_high_risk_command(context.command))
    or any(is_configuration_path(path) for path in context.affected_paths)
    or context.is_dangerous
)

if requires_confirmation:
    return ApprovalDecision.NEEDS_CONFIRMATION
""",
        "explain": "比如网络访问、删除文件、git push、修改 .env/config 都属于敏感操作。项目不会让模型直接执行，而是让审批策略决定。",
        "oral": "因为这是能真实操作本地环境的 Agent，所以安全审批非常重要。",
    },
    {
        "q": "12. never、yolo、auto-edit 这些审批策略有什么区别？",
        "idea": "它们决定工具调用遇到风险时是自动通过、要求确认，还是拒绝。",
        "code": """
# safety/approval.py
if self.approval_policy == ApprovalPolicy.YOLO:
    return ApprovalDecision.APPROVED

if self.approval_policy == ApprovalPolicy.NEVER:
    if is_safe_command(command):
        return ApprovalDecision.APPROVED
    return ApprovalDecision.REJECTED
""",
        "explain": "YOLO 基本全部放行；NEVER 只允许安全命令；默认 on-request 对高风险操作要求确认。这里也能看出安全策略是可配置的。",
        "oral": "不同用户对自动化程度的接受度不同，所以项目把审批做成策略。",
    },
    {
        "q": "13. 上下文管理为什么重要？",
        "idea": "Agent 多轮对话会积累用户消息、模型回答和工具结果，如果不管理，容易超过模型上下文窗口。",
        "code": """
# context/manager.py
def get_messages(self):
    messages = [{"role": "system", "content": system_prompt}]
    for item in self._messages:
        messages.append(item.to_dict())
    return messages

def needs_compression(self):
    return current_tokens > (context_limit * 0.8)
""",
        "explain": "ContextManager 每次请求模型前都会构造完整 messages，包括系统提示词和历史消息。当 token 接近窗口上限时，会触发压缩。",
        "oral": "它解决的是长对话记忆问题，让 Agent 能持续工作而不是很快上下文爆掉。",
    },
    {
        "q": "14. 项目上下文是怎么自动生成的？",
        "idea": "context/project.py 会扫描项目目录、语言、依赖文件和 README 摘要，生成 Project Overview。",
        "code": """
# context/project.py
root_entries = sorted(...)
extensions = Counter(path.suffix.lower() for path in files)
dependency_paths = [cwd / name for name in DEPENDENCY_FILES if ...]

lines = [
    "# Project Overview",
    f"- Root entries: ...",
    f"- Detected languages: ...",
]
""",
        "explain": "这样模型一进入项目，就能知道大致结构：这是 Python 还是 Node 项目，有没有 pyproject.toml、package.json、README 等。",
        "oral": "这相当于自动给模型一份项目简介，减少模型一上来完全不知道项目结构的问题。",
    },
    {
        "q": "15. 修改代码后为什么会自动跑测试？",
        "idea": "为了形成质量闭环：改代码、跑验证、失败后把结果反馈给模型继续修。",
        "code": """
# agent/agent.py
if (
    result.success
    and tool_call.name in {"edit", "write_file"}
    and is_source_file(changed_path)
):
    for check in discover_verification_checks(self.config.cwd):
        verification = await self.session.tool_registry.invoke("shell", ...)
""",
        "explain": "这里只在 edit/write_file 修改了源码文件后触发验证。验证命令由 context/verification.py 根据项目文件自动发现。",
        "oral": "这个设计让 Agent 不只是能改代码，还能自己检查改完有没有破坏项目。",
    },
    {
        "q": "16. 自动验证怎么知道运行什么命令？",
        "idea": "它根据项目文件判断。Python 项目看 pyproject.toml 和 tests 目录，Node 项目看 package.json scripts。",
        "code": """
# context/verification.py
if pyproject.is_file():
    if "ruff" in content:
        checks.append(VerificationCheck("lint", "uv run ruff check ."))
    if (cwd / "tests").is_dir():
        checks.append(
            VerificationCheck("unit tests", "uv run python -m unittest discover -s tests")
        )

if package_json.is_file():
    if "test" in scripts:
        checks.append(VerificationCheck("unit tests", "npm run test"))
""",
        "explain": "它没有硬编码只支持 Python，而是同时支持 Python 和 Node 的常见验证方式。不过规则还比较简单，后续可以做成项目配置可扩展。",
        "oral": "本项目我实际跑过测试，命令是 uv run python -m unittest discover -s tests，结果 25 个测试全部通过。",
    },
    {
        "q": "17. 会话保存和恢复怎么实现？",
        "idea": "PersistenceManager 把 session 快照保存成本地 JSON，恢复时重新构造 Session 并把历史消息塞回 ContextManager。",
        "code": """
# agent/persistence.py
class SessionSnapshot:
    session_id: str
    messages: list[dict[str, Any]]
    total_usage: TokenUsage

def save_session(self, snapshot):
    file_path = self.sessions_dir / f"{snapshot.session_id}.json"
    self._atomic_save(file_path, snapshot)
""",
        "explain": "保存的不只是文本，还包括 session_id、时间、turn_count、messages 和 token 用量。这样 /resume 能恢复之前的对话上下文。",
        "oral": "会话恢复的核心就是把历史 messages 重新加载回来，让模型接着之前的上下文继续工作。",
    },
    {
        "q": "18. 为什么保存文件要用原子写入？",
        "idea": "为了避免写 JSON 时程序崩溃导致会话文件损坏。",
        "code": """
# agent/persistence.py
fd, temporary_path = tempfile.mkstemp(...)
json.dump(snapshot.to_dict(), fp, indent=2)
fp.flush()
os.fsync(fp.fileno())
os.replace(temporary_path, file_path)
""",
        "explain": "先写临时文件，写完并刷盘后，再替换正式文件。这样要么保留旧文件，要么得到完整新文件，不容易出现半个 JSON。",
        "oral": "这是一个比较工程化的细节，说明项目考虑了异常情况下的数据一致性。",
    },
    {
        "q": "19. MCP 在这个项目里是什么作用？",
        "idea": "MCP 用来接入外部工具服务器，把外部能力变成 Agent 可调用工具。",
        "code": """
# tools/mcp/mcp_manager.py
await asyncio.gather(*connection_tasks, return_exceptions=True)

for tool_info in client.tools:
    mcp_tool = MCPTool(
        tool_info=tool_info,
        client=client,
        config=self.config,
        name=f"{client.name}__{tool_info.name}",
    )
    registry.register_mcp_tool(mcp_tool)
""",
        "explain": "MCPManager 会连接配置里的 MCP server，如果连接成功，就把服务端暴露的工具包装成 MCPTool，注册进 ToolRegistry。",
        "oral": "MCP 让工具系统不局限于项目内置工具，相当于给 Agent 做插件扩展。",
    },
    {
        "q": "20. 这个项目有哪些可以改进的地方？",
        "idea": "可以从 bug、类型、安全、测试和扩展性几个角度回答。",
        "code": """
# config/config.py
class ApprovalPolicy(str, Enum):
    AUTO = "auto"
    AUTO_EDIT = "auto-edut"  # 疑似应为 auto-edit

# client/response.py
class ToolCall:
    arguments: str = ""  # 实际解析后可能是 dict
""",
        "explain": "AUTO_EDIT 字符串拼写错误会影响用户通过命令设置审批策略。ToolCall.arguments 类型标注和实际使用不一致，容易造成维护困惑。此外 EditTool 可以考虑原子写入，自动验证可以支持自定义命令。",
        "oral": "如果让我优化，我会先修明显 bug，再补测试，然后增强验证配置和文件写入安全性。",
    },
    {
        "q": "21. 如果让你现场手撕核心代码，你写什么？",
        "idea": "优先写简化版 Agent Loop，因为它最能体现这个项目的核心。",
        "code": """
async def agent_loop(user_message):
    messages.append({"role": "user", "content": user_message})

    for _ in range(max_turns):
        response = await llm.chat(messages, tools=tool_schemas)

        if not response.tool_calls:
            return response.text

        messages.append({
            "role": "assistant",
            "content": response.text,
            "tool_calls": response.tool_calls,
        })

        for call in response.tool_calls:
            result = await tool_registry.invoke(call.name, call.arguments)
            messages.append({
                "role": "tool",
                "tool_call_id": call.id,
                "content": result,
            })
""",
        "explain": "这段代码保留了核心思想：用户消息进入上下文，模型选择工具，程序执行工具，工具结果再进入上下文，然后继续下一轮。",
        "oral": "现场手撕不需要完全复刻项目代码，只要把模型和工具的闭环写清楚，就抓住了核心。",
    },
    {
        "q": "22. 如果让你手撕工具注册中心，你怎么写？",
        "idea": "写一个字典保存工具名和工具实例，invoke 时按名字找到工具执行。",
        "code": """
class Tool:
    name = ""

    async def execute(self, params):
        raise NotImplementedError


class ToolRegistry:
    def __init__(self):
        self.tools = {}

    def register(self, tool):
        self.tools[tool.name] = tool

    async def invoke(self, name, params):
        if name not in self.tools:
            return {"ok": False, "error": "unknown tool"}
        return await self.tools[name].execute(params)
""",
        "explain": "真实项目在这个基础上增加了 Pydantic 参数校验、安全审批、hook、MCP 工具和异常处理。",
        "oral": "简化版体现思路，真实版补齐工程能力。",
    },
]


def build():
    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("lm-program 项目面试考核逐题详解版")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("每个问题都带回答思路、代码结合、具体解释和口语化表述")
    set_font(run, size=11, color="555555")

    doc.add_heading("使用建议", level=1)
    bullet(doc, "先背 30 秒项目介绍，再重点理解 Agent Loop、ToolRegistry、ApprovalManager、ContextManager。")
    bullet(doc, "老师追问代码时，不要只说文件名，要说清楚这段代码解决了什么问题。")
    bullet(doc, "如果要求手撕代码，优先手撕 Agent Loop 或 ToolRegistry。")

    for item in QUESTIONS:
        doc.add_heading(item["q"], level=1)
        para(doc, " " + item["idea"], "回答思路：")
        para(doc, "下面这段代码是可以拿出来讲的核心片段。", "结合代码：")
        code(doc, item["code"])
        para(doc, " " + item["explain"], "具体解释：")
        para(doc, " " + item["oral"], "考核说法：")

    doc.add_heading("最后背诵版总结", level=1)
    doc.add_paragraph(
        "这个项目的核心不是简单调用大模型 API，而是构建了一个能操作本地开发环境的终端 Agent。"
        "它用 Agent Loop 组织多轮推理，用 ToolRegistry 管理工具调用，用 ApprovalManager 控制风险，"
        "用 ContextManager 管理上下文，用 PersistenceManager 支持会话恢复，并在修改源码后自动运行测试。"
        "所以它的重点是把大模型能力和真实软件工程流程连接起来。"
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("lm-program 逐题详解版")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
